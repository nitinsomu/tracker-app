import io
from datetime import date

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.journal import JournalEntry
from app.schemas.journal import JournalEntryCreate, JournalEntryUpdate


async def get_entries(db: AsyncSession, user_id: int, start: date | None, end: date | None) -> list[JournalEntry]:
    stmt = select(JournalEntry).where(JournalEntry.user_id == user_id)
    if start:
        stmt = stmt.where(JournalEntry.date >= start)
    if end:
        stmt = stmt.where(JournalEntry.date <= end)
    stmt = stmt.order_by(JournalEntry.date.desc())
    result = await db.execute(stmt)
    return list(result.scalars().all())


async def get_entry(db: AsyncSession, user_id: int, entry_id: int) -> JournalEntry | None:
    result = await db.execute(
        select(JournalEntry).where(JournalEntry.id == entry_id, JournalEntry.user_id == user_id)
    )
    return result.scalar_one_or_none()


async def create_entry(db: AsyncSession, user_id: int, data: JournalEntryCreate) -> JournalEntry:
    entry = JournalEntry(user_id=user_id, **data.model_dump())
    db.add(entry)
    await db.commit()
    await db.refresh(entry)
    return entry


async def update_entry(db: AsyncSession, entry: JournalEntry, data: JournalEntryUpdate) -> JournalEntry:
    for field, value in data.model_dump(exclude_none=True).items():
        setattr(entry, field, value)
    await db.commit()
    await db.refresh(entry)
    return entry


async def delete_entry(db: AsyncSession, entry: JournalEntry) -> None:
    await db.delete(entry)
    await db.commit()


async def export_docx(entries: list[JournalEntry]) -> bytes:
    doc = Document()
    doc.add_heading("Journal Export", 0)
    for entry in sorted(entries, key=lambda e: e.date):
        doc.add_heading(str(entry.date), level=2)
        doc.add_paragraph(entry.content)
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
