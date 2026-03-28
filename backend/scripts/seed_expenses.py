"""
Seed script: parse _expense tracker.docx and insert into expenses.

Usage:
    python scripts/seed_expenses.py --file "path/to/_expense tracker.docx" --user-id 1

Line format expected (all other lines are ignored):
    DD/MM/YY - amount category, amount category, ...

Multiple expenses per line are split into individual rows.
Works with or without commas between items (e.g. "790 food 194 entertainment").
"""

import argparse
import asyncio
import re
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from sqlalchemy import select

from app.config import settings
from app.database import AsyncSessionLocal
from app.models.expense import Expense
from app.models.user import User  # noqa: F401

LINE_PATTERN = re.compile(r"^(\d{2}/\d{2}/\d{2})\s*-\s*(.+)$")
ITEM_PATTERN = re.compile(r"(\d+(?:\.\d+)?)\s+([a-zA-Z]+)", re.IGNORECASE)


def parse_line(raw: str) -> list[dict] | None:
    """Return a list of expense dicts for the line, or None if it doesn't match."""
    line = " ".join(raw.split())  # normalise whitespace
    m = LINE_PATTERN.match(line)
    if not m:
        return None

    date_str, expenses_str = m.group(1), m.group(2)

    try:
        entry_date = datetime.strptime(date_str, "%d/%m/%y").date()
    except ValueError:
        return None

    items = ITEM_PATTERN.findall(expenses_str)
    if not items:
        return None

    return [
        {
            "date": entry_date,
            "amount": float(amount),
            "category": category.strip().lower(),
        }
        for amount, category in items
    ]


def extract_lines(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines


async def seed(docx_path: str, user_id: int, dry_run: bool) -> None:
    lines = extract_lines(docx_path)

    all_entries: list[dict] = []
    matched_lines = 0
    skipped_lines = []
    for line in lines:
        result = parse_line(line)
        if result is not None:
            matched_lines += 1
            all_entries.extend(result)
        else:
            skipped_lines.append(line)

    print(f"Lines in document  : {len(lines)}")
    print(f"Lines matched      : {matched_lines}")
    print(f"Lines skipped      : {len(skipped_lines)}")
    print(f"Expense rows parsed: {len(all_entries)}")

    if not all_entries:
        print("Nothing to insert.")
        return

    if dry_run:
        print("\n-- DRY RUN (first 5 matched entries) --")
        for e in all_entries[:5]:
            print(e)
        if skipped_lines:
            print("\n-- SKIPPED LINES --")
            for l in skipped_lines:
                print(f"  {l}")
        return

    async with AsyncSessionLocal() as db:
        result = await db.execute(select(User).where(User.id == user_id))
        if not result.scalar_one_or_none():
            print(f"Error: no user with id={user_id} found.")
            return

        inserted = 0
        for e in all_entries:
            db.add(Expense(user_id=user_id, **e))
            inserted += 1

        await db.commit()

    print(f"\nDone. Inserted: {inserted} expense rows.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Seed expenses from a Word doc.")
    parser.add_argument("--file", required=True, help="Path to the .docx file")
    parser.add_argument("--user-id", required=True, type=int, help="User ID to assign entries to")
    parser.add_argument("--dry-run", action="store_true", help="Parse only, do not insert")
    args = parser.parse_args()

    asyncio.run(seed(args.file, args.user_id, args.dry_run))


if __name__ == "__main__":
    main()
