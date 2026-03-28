"""
Seed script: parse Daily New Thing.docx and insert into journal_entries.

Usage:
    python scripts/seed_journal.py --file "path/to/Daily New Thing.docx" --user-id 1

Line format expected (all other lines are ignored):
    DD/MM/YY - any free text content
"""

import argparse
import asyncio
import re
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from sqlalchemy import select

from app.database import AsyncSessionLocal
from app.models.journal import JournalEntry
from app.models.user import User  # noqa: F401

LINE_PATTERN = re.compile(r"^(\d{2}/\d{2}/\d{2})\s*-\s*(.+)$")


def parse_line(raw: str) -> dict | None:
    """Return parsed dict or None if the line doesn't match the expected format."""
    line = " ".join(raw.split())  # normalise whitespace
    m = LINE_PATTERN.match(line)
    if not m:
        return None

    date_str, content = m.group(1), m.group(2).strip()

    try:
        entry_date = datetime.strptime(date_str, "%d/%m/%y").date()
    except ValueError:
        return None

    return {"date": entry_date, "content": content}


def extract_lines(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines


async def seed(docx_path: str, user_id: int, dry_run: bool) -> None:
    lines = extract_lines(docx_path)
    parsed = [(l, parse_line(l)) for l in lines]
    skipped_lines = [l for l, e in parsed if e is None]
    entries = [e for _, e in parsed if e is not None]

    print(f"Lines in document : {len(lines)}")
    print(f"Lines matched      : {len(entries)}")
    print(f"Lines skipped      : {len(skipped_lines)}")

    if not entries:
        print("Nothing to insert.")
        return

    if dry_run:
        print("\n-- DRY RUN (first 5 matched entries) --")
        for e in entries[:5]:
            print(e)
        if skipped_lines:
            print("\n-- SKIPPED LINES --")
            for l in skipped_lines:
                print(f"  {l}")
        return

    async with AsyncSessionLocal() as db:
        result = await db.execute(select(User).where(User.id == user_id))
        if not result.scalar_one_or_none():
            print(f"Error: no user with id={user_id} found.")
            return

        inserted = skipped = 0
        for e in entries:
            exists = await db.execute(
                select(JournalEntry).where(
                    JournalEntry.user_id == user_id,
                    JournalEntry.date == e["date"],
                )
            )
            if exists.scalar_one_or_none():
                skipped += 1
                continue

            db.add(JournalEntry(user_id=user_id, **e))
            inserted += 1

        await db.commit()

    print(f"\nDone. Inserted: {inserted}  |  Skipped (duplicate date): {skipped}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Seed journal_entries from a Word doc.")
    parser.add_argument("--file", required=True, help="Path to the .docx file")
    parser.add_argument("--user-id", required=True, type=int, help="User ID to assign entries to")
    parser.add_argument("--dry-run", action="store_true", help="Parse only, do not insert")
    args = parser.parse_args()

    asyncio.run(seed(args.file, args.user_id, args.dry_run))


if __name__ == "__main__":
    main()
